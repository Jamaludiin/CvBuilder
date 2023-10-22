from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Word document
document = Document()

# Add a heading to the document
heading = document.add_heading('Curriculum Vitae', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add personal information section
personal_info_heading = document.add_heading('Personal Information', level=1)
personal_info_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
table = document.add_table(rows=4, cols=2)
table.style = 'Table Grid'
row_cells = table.rows[0].cells
row_cells[0].text = 'Full Name'
row_cells[1].text = 'John Doe'
row_cells = table.rows[1].cells
row_cells[0].text = 'Email'
row_cells[1].text = 'johndoe@example.com'
row_cells = table.rows[2].cells
row_cells[0].text = 'Phone'
row_cells[1].text = '555-123-4567'
row_cells = table.rows[3].cells
row_cells[0].text = 'Address'
row_cells[1].text = '123 Main St, Anytown USA'

# Add education section
education_heading = document.add_heading('Education', level=1)
education_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph()
p.add_run('Bachelor of Science in Computer Science').bold = True
p.add_run('\nUniversity of Anytown, USA, 2010-2014')
p = document.add_paragraph()
p.add_run('Master of Science in Computer Science').bold = True
p.add_run('\nUniversity of Anytown, USA, 2014-2016')

# Add work experience section
work_experience_heading = document.add_heading('Work Experience', level=1)
work_experience_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph()
p.add_run('Software Engineer').bold = True
p.add_run('\nABC Company, Anytown USA, 2016-present')
p = document.add_paragraph('Responsibilities:')
p.add_run('\n- Developed and maintained software applications')
p.add_run('\n- Collaborated with cross-functional teams to design and implement new features')
p.add_run('\n- Troubleshot and resolved software issues')

# Add skills section
skills_heading = document.add_heading('Skills', level=1)
skills_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph()
p.add_run('Programming Languages: ').bold = True
p.add_run('Python, Java, C++')
p = document.add_paragraph()
p.add_run('Tools and Frameworks: ').bold = True
p.add_run('Django, Flask, Git, JIRA')

# Add a photo to the document
document.add_picture('/Users/jamalabdullahi/Python Tutorial/CV BUILDERS/profile.jpeg', width=Inches(1.25))

# Save the document
document.save('cv_template2.docx')






# NOTHER CODE BUT YOU HAVE TO COMMENT THE ABOVE CODE
"""from docx import Document
from docx.shared import Inches

# create a new Word document
document = Document('/Users/jamalabdullahi/Python Tutorial/CV BUILDERS/Professional Resume.docx')

# add personal information section
document.add_heading('Personal Information', level=1)
table = document.add_table(rows=4, cols=2)
table.style = 'Table Grid'
row_cells = table.rows[0].cells
row_cells[0].text = 'Name'
row_cells[1].text = 'John Doe'
row_cells = table.rows[1].cells
row_cells[0].text = 'Address'
row_cells[1].text = '123 Main St, Anytown USA'
row_cells = table.rows[2].cells
row_cells[0].text = 'Phone'
row_cells[1].text = '555-123-4567'
row_cells = table.rows[3].cells
row_cells[0].text = 'Email'
row_cells[1].text = 'johndoe@example.com'

# add education section
document.add_heading('Education', level=1)
p = document.add_paragraph()
p.add_run('Bachelor of Science in Computer Science').bold = True
p.add_run('\nUniversity of Anytown, USA, 2010-2014')
p = document.add_paragraph()
p.add_run('Master of Science in Computer Science').bold = True
p.add_run('\nUniversity of Anytown, USA, 2014-2016')

# add work experience section
document.add_heading('Work Experience', level=1)
p = document.add_paragraph()
p.add_run('Software Engineer').bold = True
p.add_run('\nABC Company, Anytown USA, 2016-present')
p = document.add_paragraph('Responsibilities:')
p.add_run('\n- Developed and maintained software applications')
p.add_run('\n- Collaborated with cross-functional teams to design and implement new features')
p.add_run('\n- Troubleshot and resolved software issues')

# add skills section
document.add_heading('Skills', level=1)
p = document.add_paragraph()
p.add_run('Programming Languages: ').bold = True
p.add_run('Python, Java, C++')
p = document.add_paragraph()
p.add_run('Tools and Frameworks: ').bold = True
p.add_run('Django, Flask, Git, JIRA')

# add a photo to the document
document.add_picture('/Users/jamalabdullahi/Python Tutorial/CV BUILDERS/profile.jpeg', width=Inches(1.25))

# save the document
document.save('cv_template3.docx')
"""