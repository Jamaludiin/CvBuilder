from docx import Document
from docx.shared import Inches

# Create a new document
document = Document()

# Add a heading to the document
document.add_heading('Curriculum Vitae', 0)

# Prompt the user for their personal information
name = input('Enter your name: ')
email = input('Enter your email address: ')
phone = input('Enter your phone number: ')
address = input('Enter your address: ')

# Add the personal information to the document
document.add_heading('Personal Information', level=1)
document.add_paragraph(f'Name: {name}')
document.add_paragraph(f'Email: {email}')
document.add_paragraph(f'Phone: {phone}')
document.add_paragraph(f'Address: {address}')

# Prompt the user for their education information
document.add_heading('Education', level=1)

while True:
    degree = input('Enter your degree (type "q" to quit): ')
    if degree == 'q':
        break
    field_of_study = input('Enter your field of study: ')
    institution = input('Enter the name of the institution: ')
    location = input('Enter the location of the institution: ')
    start_date = input('Enter the start date: ')
    end_date = input('Enter the end date: ')

    # Add the education information to the document
    p = document.add_paragraph()
    p.add_run(f'{degree}, {field_of_study}').bold = True
    p.add_run(f'\n{institution}, {location}')
    p.add_run(f'\n{start_date} - {end_date}\n').italic = True

# Prompt the user for their work experience information
document.add_heading('Work Experience', level=1)

while True:
    position = input('Enter your position (type "q" to quit): ')
    if position == 'q':
        break
    employer = input('Enter the name of the employer: ')
    location = input('Enter the location of the employer: ')
    start_date = input('Enter the start date: ')
    end_date = input('Enter the end date: ')
    description = input('Enter a brief description of your work: ')

    # Add the work experience information to the document
    p = document.add_paragraph()
    p.add_run(f'{position} - {employer}').bold = True
    p.add_run(f'\n{location}')
    p.add_run(f'\n{start_date} - {end_date}\n').italic = True
    p.add_run(description)

# Prompt the user for their skills information
document.add_heading('Skills', level=1)
skills = input('Enter your skills, separated by commas: ')

# Add the skills information to the document
document.add_paragraph(skills)

# Save the document
document.save('cv.docx')
print('Your CV has been saved to cv.docx.')
